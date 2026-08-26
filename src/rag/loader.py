"""文档加载与切片模块

支持 Word（.docx）和 Markdown（.md）格式文档的加载与语义段落切片。
"""

from pathlib import Path

import docx
import markdown
from llama_index.core import Document


def load_docx(file_path: Path) -> str:
    """加载 Word 文档，提取纯文本内容"""
    doc = docx.Document(str(file_path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return "\n\n".join(paragraphs)


def load_markdown(file_path: Path) -> str:
    """加载 Markdown 文档，提取纯文本内容"""
    raw = file_path.read_text(encoding="utf-8")
    # 将 markdown 转为纯文本，保留段落结构
    html = markdown.markdown(raw)
    # 简单去除 HTML 标签
    import re
    text = re.sub(r"<[^>]+>", "", html)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def load_document(file_path: Path) -> Document:
    """加载单个文档，返回 LlamaIndex Document 对象"""
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        text = load_docx(file_path)
    elif suffix in (".md", ".markdown"):
        text = load_markdown(file_path)
    elif suffix == ".txt":
        text = file_path.read_text(encoding="utf-8")
    else:
        raise ValueError(f"不支持的文档格式: {suffix}")

    return Document(
        text=text,
        metadata={
            "file_name": file_path.name,
            "file_path": str(file_path),
            "file_type": suffix,
        },
    )


def load_documents(directory: Path) -> list[Document]:
    """批量加载目录下的所有文档"""
    supported_extensions = {".docx", ".md", ".markdown", ".txt"}
    documents = []

    if not directory.exists():
        return documents

    for file_path in sorted(directory.rglob("*")):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            try:
                doc = load_document(file_path)
                documents.append(doc)
            except Exception as e:
                print(f"⚠️ 加载文档失败 {file_path}: {e}")

    return documents
